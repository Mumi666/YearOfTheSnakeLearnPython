from docx import Document
from docx.shared import RGBColor


def merge_doc(file_list: list, output_path: str = './result/new.docx'):
    """
    合并多个Word文档

    Args:
        file_list: 要合并的Word文档文件路径列表
        output_path: 输出文件路径，默认为当前目录下的new.docx
    """
    # 创建新的Document对象
    doc = Document()

    for file in sorted(file_list):
        another_file = Document(file)
        for para in another_file.paragraphs:
            new_para = doc.add_paragraph('')
            new_para.add_run(para.text)

    # 保存合并后的文档
    doc.save(output_path)
    print(f"文档已成功合并并保存到: {output_path}")

def add_content_model(content):
    doc = Document()
    para = doc.add_paragraph('').add_run(content)
    para.underline = True
    para.font.bold = True
    para.font.color.rgb = RGBColor(255,128,128)
